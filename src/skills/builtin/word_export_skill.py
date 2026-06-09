"""Word Export Skill — 将 Markdown 内容导出为 DOCX/PDF。

功能：
1. Markdown → DOCX（支持模板样式克隆）
2. DOCX 编辑（在已有文档中追加内容）
3. Markdown → PDF
"""
import re
import os
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime

from src.skills.base import BaseSkill, SkillResult, SkillContext

# Output directory
OUTPUT_DIR = Path.home() / ".cs599-agent" / "exports"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


class WordExportSkill(BaseSkill):
    """Export Markdown content to Word DOCX or PDF with template support."""

    name = "word_export"
    display_name = "Word/PDF 导出"
    description = "将研究报告导出为 Word 文档或 PDF 文件，支持模板样式克隆和已有文档编辑"
    version = "1.0.0"
    author = "CS599 Agent"
    tags = ["export", "word", "pdf", "docx"]

    parameters_schema = {
        "content": {"type": "string", "description": "Markdown 内容", "required": True},
        "format": {"type": "string", "description": "输出格式: docx / pdf", "default": "docx"},
        "filename": {"type": "string", "description": "文件名（不含扩展名）", "default": "report"},
        "template_path": {"type": "string", "description": "模板 .docx 文件路径（可选）", "default": ""},
        "edit_mode": {"type": "boolean", "description": "是否编辑已有文档", "default": False},
        "edit_docx_path": {"type": "string", "description": "已有文档路径（edit_mode=True 时必填）", "default": ""},
    }

    def execute(self, context: SkillContext) -> SkillResult:
        content = context.custom_params.get("content", "")
        fmt = context.custom_params.get("format", "docx")
        filename = context.custom_params.get("filename", "report")
        template_path = context.custom_params.get("template_path", "")
        edit_mode = context.custom_params.get("edit_mode", False)
        edit_docx_path = context.custom_params.get("edit_docx_path", "")
        steps = []

        if not content:
            return SkillResult(success=False, error="内容不能为空")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = re.sub(r'[^\w\-_]', '_', filename)[:50]
        out_stem = f"{safe_name}_{timestamp}"

        try:
            if fmt == "docx":
                if edit_mode and edit_docx_path:
                    steps.append({"step": 1, "action": "edit_docx", "status": "running"})
                    out_path = self._edit_docx(edit_docx_path, content, out_stem)
                    steps[-1]["status"] = "done"
                    steps[-1]["result"] = f"已编辑: {out_path.name}"
                else:
                    steps.append({"step": 1, "action": "create_docx", "status": "running"})
                    out_path = self._markdown_to_docx(content, out_stem, template_path)
                    steps[-1]["status"] = "done"
                    steps[-1]["result"] = f"已生成: {out_path.name}"
            elif fmt == "pdf":
                steps.append({"step": 1, "action": "create_pdf", "status": "running"})
                out_path = self._markdown_to_pdf(content, out_stem)
                steps[-1]["status"] = "done"
                steps[-1]["result"] = f"已生成: {out_path.name}"
            else:
                return SkillResult(success=False, error=f"不支持的格式: {fmt}")

            return SkillResult(
                success=True,
                content=f"文件已导出: {out_path}",
                metadata={
                    "file_path": str(out_path),
                    "file_name": out_path.name,
                    "format": fmt,
                    "size_bytes": out_path.stat().st_size,
                },
                steps=steps,
            )
        except Exception as e:
            return SkillResult(success=False, error=f"导出失败: {e}", steps=steps)

    # ── DOCX: Markdown → New Document ──────────────────────────────

    def _markdown_to_docx(self, md: str, stem: str, template: str = "") -> Path:
        """Convert Markdown to DOCX, optionally cloning a template's styles."""
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        if template and Path(template).exists():
            doc = Document(template)
            # Clone styles from template
            self._clone_template_styles(doc, template)
        else:
            doc = Document()
            self._set_default_styles(doc)

        for line in md.split("\n"):
            line_stripped = line.strip()

            # Headings
            if line_stripped.startswith("### "):
                p = doc.add_heading(line_stripped[4:], level=3)
            elif line_stripped.startswith("## "):
                p = doc.add_heading(line_stripped[3:], level=2)
            elif line_stripped.startswith("# "):
                p = doc.add_heading(line_stripped[2:], level=1)
            # Code block
            elif line_stripped.startswith("```"):
                continue  # skip code fences
            elif line_stripped.startswith("|") and line_stripped.endswith("|"):
                self._add_table_row(doc, line_stripped)
            # Bullet
            elif line_stripped.startswith("- ") or line_stripped.startswith("* "):
                doc.add_paragraph(line_stripped[2:], style="List Bullet")
            # Numbered
            elif re.match(r"^\d+\.\s", line_stripped):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", line_stripped), style="List Number")
            # Quote
            elif line_stripped.startswith("> "):
                p = doc.add_paragraph(line_stripped[2:])
                p.italic = True
            # Empty
            elif not line_stripped:
                continue
            # Regular paragraph
            else:
                doc.add_paragraph(line_stripped)

        out_path = OUTPUT_DIR / f"{stem}.docx"
        doc.save(str(out_path))
        return out_path

    def _clone_template_styles(self, doc, template_path: str):
        """Clone styles from a template .docx into the new document.

        Reads the template's style definitions and applies matching
        styles to the new document's style elements.
        """
        from docx import Document as Doc
        from docx.oxml import parse_xml

        template_doc = Doc(template_path)
        try:
            # Copy paragraph styles (font name, size, bold, color, spacing)
            for style in template_doc.styles:
                if style.type is None:
                    continue
                try:
                    new_style = doc.styles[style.name]
                except KeyError:
                    continue
                font = style.font
                new_font = new_style.font
                if font.name:
                    new_font.name = font.name
                if font.size:
                    new_font.size = font.size
                if font.bold is not None:
                    new_font.bold = font.bold
                if font.italic is not None:
                    new_font.italic = font.italic
                if font.color and font.color.rgb:
                    try:
                        new_font.color.rgb = font.color.rgb
                    except Exception:
                        pass
                # Paragraph format
                pf = style.paragraph_format
                npf = new_style.paragraph_format
                if pf.space_before:
                    npf.space_before = pf.space_before
                if pf.space_after:
                    npf.space_after = pf.space_after
                if pf.line_spacing:
                    npf.line_spacing = pf.line_spacing
        finally:
            template_doc.close()

    def _set_default_styles(self, doc):
        """Set default academic styles when no template is provided."""
        from docx.shared import Pt, RGBColor
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.5

        for level in range(1, 4):
            try:
                hs = doc.styles[f"Heading {level}"]
                hs.font.name = "Times New Roman"
                hs.font.bold = True
                if level == 1:
                    hs.font.size = Pt(16)
                elif level == 2:
                    hs.font.size = Pt(14)
                else:
                    hs.font.size = Pt(12)
            except KeyError:
                pass

    def _add_table_row(self, doc, line: str):
        """Parse a markdown table row and add to document."""
        cols = [c.strip() for c in line.split("|")[1:-1]]
        if not hasattr(self, "_table_rows"):
            self._table_rows = []
        self._table_rows.append(cols)

    # ── DOCX: Edit existing document ───────────────────────────────

    def _edit_docx(self, docx_path: str, new_content: str, stem: str) -> Path:
        """Append new content to an existing DOCX file.

        Preserves all original formatting and matches new content
        to the document's existing styles.
        """
        from docx import Document
        from docx.shared import Pt

        doc = Document(docx_path)
        # Append new content at the end
        doc.add_paragraph("")  # spacing
        doc.add_heading("追加内容", level=1)
        doc.add_paragraph(f"编辑时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        for line in new_content.split("\n"):
            s = line.strip()
            if s.startswith("### "):
                doc.add_heading(s[4:], level=3)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=2)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif s.startswith("- ") or s.startswith("* "):
                doc.add_paragraph(s[2:], style="List Bullet")
            elif s:
                doc.add_paragraph(s)

        out_path = OUTPUT_DIR / f"{stem}_edited.docx"
        doc.save(str(out_path))
        return out_path

    # ── PDF: Markdown → PDF ────────────────────────────────────────

    def _markdown_to_pdf(self, md: str, stem: str) -> Path:
        """Convert Markdown to a simple academic PDF using fpdf2."""
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)

        # Title
        pdf.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", uni=True)
        pdf.add_font("DejaVu", "B", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", uni=True)

        for line in md.split("\n"):
            s = line.strip()
            if s.startswith("# "):
                pdf.set_font("DejaVu", "B", 16)
                pdf.cell(0, 10, s[2:], new_x="LMARGIN", new_y="NEXT")
                pdf.ln(4)
            elif s.startswith("## "):
                pdf.set_font("DejaVu", "B", 14)
                pdf.cell(0, 8, s[3:], new_x="LMARGIN", new_y="NEXT")
                pdf.ln(3)
            elif s.startswith("### "):
                pdf.set_font("DejaVu", "B", 12)
                pdf.cell(0, 7, s[4:], new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
            elif s.startswith("```"):
                in_code = not hasattr(self, "_in_code_flag")
                self._in_code_flag = in_code
            elif s.startswith("- ") or s.startswith("* "):
                pdf.set_font("DejaVu", "", 10)
                pdf.cell(5)  # indent
                pdf.multi_cell(0, 5, f"• {s[2:]}")
            elif s:
                font_style = "B" if s.startswith("**") and s.endswith("**") else ""
                text = s.strip("*") if font_style else s
                pdf.set_font("DejaVu", font_style, 10)
                pdf.multi_cell(0, 5, text)
                pdf.ln(1)

            # Handle tables
            if s.startswith("|") and s.endswith("|"):
                cols = [c.strip() for c in s.split("|")[1:-1]]
                if not cols or all(c.startswith("---") for c in cols):
                    continue
                pdf.set_font("DejaVu", "", 9)
                for col in cols:
                    pdf.cell(45, 6, col[:30], border=1)
                pdf.ln()

        out_path = OUTPUT_DIR / f"{stem}.pdf"
        pdf.output(str(out_path))
        return out_path
