"""Office Export Skill — Convert research reports to Word (.docx) and PDF.

References the report structure from Codex Academic Skills (office-academic-skill).
Default language: English; also supports Chinese.
"""
import re
import time
from pathlib import Path
from typing import List, Dict
from datetime import datetime

from src.skills.base import BaseSkill, SkillResult, SkillContext

# Output directory
OUTPUT_DIR = Path.home() / ".cs599-agent" / "output"


# ── DOCX generation ──────────────────────────────────────────────────────────
def _render_docx(markdown_text: str, title: str, author: str, language: str) -> Path:
    """Convert Markdown to .docx using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Default style ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── Helper: add a styled heading ──
    def add_heading(text: str, level: int):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'Times New Roman'
        return h

    # ── Title page ──
    for _ in range(6):
        doc.add_paragraph('')
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(title or "Research Report")
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(30, 30, 80)

    doc.add_paragraph('')
    if author:
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = ap.add_run(f"Author: {author}")
        ar.font.size = Pt(14)
        ar.font.name = 'Times New Roman'
        ar.font.color.rgb = RGBColor(60, 60, 60)

    doc.add_paragraph('')
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(datetime.now().strftime("%B %d, %Y"))
    dr.font.size = Pt(12)
    dr.font.name = 'Times New Roman'
    dr.font.color.rgb = RGBColor(100, 100, 100)

    # Generate abstract (first 300 chars of content as summary)
    plain = re.sub(r'[#*>|`\-\[\]]', '', markdown_text)[:300]
    if plain:
        doc.add_paragraph('')
        ab = doc.add_paragraph()
        ab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ab_lab = ab.add_run("Abstract")
        ab_lab.bold = True
        ab_lab.font.size = Pt(12)
        ab_lab.font.name = 'Times New Roman'
        doc.add_paragraph('')
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ap_r = ap.add_run(plain + "...")
        ap_r.font.size = Pt(10)
        ap_r.font.name = 'Times New Roman'
        ap_r.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_page_break()

    # ── Markdown parser (same as before, but uses structured headings) ──
    dr.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ── Markdown parser ──
    lines = markdown_text.split('\n')
    i = 0
    in_code = False
    code_buf: List[str] = []
    in_table = False
    table_buf: List[str] = []

    def flush_code():
        if not code_buf:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        pf = p.paragraph_format
        pf.left_indent = Inches(0.3)
        # Add shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F0F0F0')
        shading.set(qn('w:val'), 'clear')
        p.paragraph_format.element.get_or_add_pPr().append(shading)
        for cl in code_buf:
            r = p.add_run(cl + '\n')
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(50, 50, 50)
        code_buf.clear()

    def flush_table():
        if len(table_buf) < 2:
            table_buf.clear()
            return
        rows_data = []
        for line in table_buf:
            if line.strip().startswith('|') and line.strip().endswith('|'):
                cols = [c.strip() for c in line.split('|')[1:-1]]
                rows_data.append(cols)
        if not rows_data:
            table_buf.clear()
            return
        # Determine columns
        max_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data)-1, cols=max_cols)
        table.style = 'Table Grid'
        # Header
        header = rows_data[0]
        for j, h in enumerate(header):
            cell = table.rows[0].cells[j]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
            # Shading for header
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9E2F3')
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)
        # Data rows (skip separator line and header)
        row_idx = 1
        for ri in range(2, len(rows_data)):
            if ri < len(rows_data) and row_idx < len(table.rows):
                for j, val in enumerate(rows_data[ri]):
                    if j < len(table.rows[row_idx].cells):
                        table.rows[row_idx].cells[j].text = val
                row_idx += 1
        table_buf.clear()

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            if in_code:
                in_code = False
                flush_code()
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Flush table if we encounter a non-table line
        if in_table and not (line.strip().startswith('|') and line.strip().endswith('|')):
            in_table = False
            flush_table()

        # Table
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # Skip separator line (|---|---|)
            if not re.match(r'^\|[\s\-:]+\|', line):
                in_table = True
                table_buf.append(line)
            i += 1
            continue

        # Heading 1
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:], level=1)
        # Heading 2
        elif line.startswith('## ') and not line.startswith('### '):
            doc.add_heading(line[3:], level=2)
        # Heading 3
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # Blockquote
        elif line.startswith('> '):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.italic = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(80, 80, 80)
            pf = p.paragraph_format
            pf.left_indent = Inches(0.3)
            # Add left border effect
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:color'), '888888')
            pBdr.append(left)
            pPr.append(pBdr)
        # List item
        elif re.match(r'^[\-\*]\s', line):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
        # Empty line
        elif line.strip() == '':
            pass
        # Regular paragraph
        else:
            p = doc.add_paragraph(line)

        i += 1

    if in_code:
        flush_code()
    if in_table:
        flush_table()

    # ── Save ──
    safe_title = re.sub(r'[^\w\s\-]', '', title)[:40] or "report"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / f"{safe_title}_{ts}.docx"
    doc.save(str(path))
    return path


# ── PDF generation ───────────────────────────────────────────────────────────
def _render_pdf(markdown_text: str, title: str, author: str, language: str) -> Path:
    """Convert Markdown to PDF using fpdf2."""
    from fpdf import FPDF

    class AcademicPDF(FPDF):
        def __init__(self):
            super().__init__()
            self.title_short = title[:60] if title else "Report"

        def header(self):
            if self.page_no() > 1:
                self.set_font('Times', 'I', 8)
                self.set_text_color(150, 150, 150)
                self.cell(0, 8, self.title_short, align='R')
                self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font('Times', 'I', 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, str(self.page_no()), align='C')

    pdf = AcademicPDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Try to add Chinese font (NotoSansSC) if available
    font_paths = [
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
        Path.home() / ".fonts" / "NotoSansSC-Regular.ttf",
    ]
    has_cn_font = False
    cn_font_name = 'CN'
    for fp in font_paths:
        if fp.exists():
            pdf.add_font(cn_font_name, '', str(fp), uni=True)
            has_cn_font = True
            break

    def write_h(text: str, size: int):
        """Write heading - use larger size, no bold style to avoid fpdf2 font derivation issues."""
        pdf.set_font('Times', 'B', size)
        pdf.multi_cell(0, size * 0.5, text)
        pdf.ln(2)

    def write_p(text: str, size: int = 11):
        """Write paragraph."""
        pdf.set_font('Times', '', size)
        pdf.multi_cell(0, 5, text)

    # Title page
    pdf.add_page()
    pdf.ln(40)
    pdf.set_font('Times', 'B', 24)
    pdf.multi_cell(0, 12, title or "Research Report", align='C')
    pdf.ln(10)
    if author:
        pdf.set_font('Times', '', 14)
        pdf.cell(0, 10, f"Author: {author}", align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.add_page()
    pdf.ln(40)
    pdf.set_font('Times', 'B', 24)
    pdf.multi_cell(0, 12, title or "Research Report", align='C')
    pdf.ln(10)
    if author:
        pdf.set_font('Times', '', 14)
        pdf.cell(0, 10, f"Author: {author}", align='C', new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)
    pdf.set_font('Times', '', 11)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 10, datetime.now().strftime("%B %d, %Y"), align='C')
    pdf.set_text_color(0, 0, 0)
    pdf.add_page()

    # Markdown parser
    lines = markdown_text.split('\n')
    i = 0
    in_code = False
    code_buf: List[str] = []

    def get_font():
        return 'CN' if has_cn_font else 'Times'

    def write_code():
        if not code_buf:
            return
        pdf.set_fill_color(245, 245, 245)
        pdf.set_font('Courier', '', 8)
        for cl in code_buf:
            pdf.multi_cell(0, 4.5, cl)
        pdf.ln(3)
        code_buf.clear()

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('```'):
            if in_code:
                in_code = False
                write_code()
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Skip table separator lines (too complex for basic PDF, render as text)
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if re.match(r'^\|[\s\-:]+\|', line):
                i += 1
                continue
            # Render table line as text with tabs
            cols = [c.strip() for c in line.split('|')[1:-1]]
            pdf.set_font('Times', '', 9)
            pdf.cell(0, 6, ' | '.join(cols))
            pdf.ln()
            i += 1
            continue

        if line.startswith('# ') and not line.startswith('## '):
            write_h(line[2:], 16)
        elif line.startswith('## ') and not line.startswith('### '):
            write_h(line[3:], 14)
        elif line.startswith('### '):
            write_h(line[4:], 12)
        elif line.startswith('> '):
            pdf.set_font('Times', 'I', 10)
            pdf.set_x(pdf.get_x() + 8)
            pdf.multi_cell(0, 5, line[2:], new_x="LMARGIN")
            pdf.ln(1)
        elif line.strip() == '':
            pdf.ln(3)
        else:
            write_p(line)
        i += 1

    if in_code:
        write_code()

    safe_title = re.sub(r'[^\w\s\-]', '', title)[:40] or "report"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / f"{safe_title}_{ts}.pdf"
    pdf.output(str(path))
    return path


# ── Skill class ──────────────────────────────────────────────────────────────
class OfficeExportSkill(BaseSkill):
    """Export markdown content to Word (.docx) or PDF.

    References the report structure from Codex office-academic-skill.
    Default language: English; supports Chinese via language='zh'.
    """
    name = "office_export"
    display_name = "学术文档导出"
    description = "将研究报告/论文导出为 Word (.docx) 或 PDF 格式。支持英文和中文。"
    version = "1.0.0"
    author = "CS599 Agent"
    tags = ["export", "office", "word", "pdf"]

    parameters_schema = {
        "content": {"type": "string", "description": "要导出的 Markdown 内容", "required": True},
        "format": {"type": "string", "description": "输出格式", "options": ["docx", "pdf"], "default": "docx"},
        "title": {"type": "string", "description": "文档标题", "default": "Research Report"},
        "author": {"type": "string", "description": "作者名", "default": "CS599 Research Assistant"},
        "language": {"type": "string", "description": "语言", "options": ["en", "zh"], "default": "en"},
    }

    def execute(self, context: SkillContext) -> SkillResult:
        content = context.custom_params.get("content", "")
        fmt = context.custom_params.get("format", "docx")
        title = context.custom_params.get("title", "Research Report")
        author = context.custom_params.get("author", "CS599 Research Assistant")
        language = context.custom_params.get("language", "en")

        if not content:
            return SkillResult(success=False, error="No content provided")

        # If content is short (likely not a full report), wrap it
        if len(content) < 100:
            content = f"# {title}\n\n{content}"

        try:
            if fmt == "pdf":
                path = _render_pdf(content, title, author, language)
            else:
                path = _render_docx(content, title, author, language)

            filename = path.name
            download_url = f"/api/download/{filename}"

            return SkillResult(
                success=True,
                content=download_url,
                metadata={
                    "filename": filename,
                    "format": fmt,
                    "path": str(path),
                    "size": path.stat().st_size,
                },
            )
        except ImportError as e:
            missing = str(e).split("'")[1] if "'" in str(e) else str(e)
            return SkillResult(
                success=False,
                error=f"缺少依赖库: {missing}。请运行: pip install python-docx fpdf2",
            )
        except Exception as e:
            return SkillResult(success=False, error=f"文档导出失败: {e}")